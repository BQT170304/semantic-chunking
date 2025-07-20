from __future__ import annotations

import os
from io import BytesIO
from typing import List
from typing import Tuple

import cv2
import docx
import numpy as np
import openpyxl
import pandas as pd
import pdfplumber
import pytesseract
from fastapi import UploadFile
from pdf2image import convert_from_bytes
from pdfplumber.utils import extract_text
from pdfplumber.utils import get_bbox_overlap
from pdfplumber.utils import obj_to_bbox
from PIL import Image
from PIL import ImageDraw
from PIL import ImageEnhance

from .base import FileType


class ExtractorService:
    """Service for extracting raw text from various file formats."""

    def extract(self, file: UploadFile) -> str:
        """Extract raw text from the given file based on its format."""
        file_type = self.get_file_type(file)

        extraction_methods = {
            FileType.PDF: self.extract_pdf,
            FileType.DOCX: self.extract_docx,
            FileType.XLSX: self.extract_xlsx,
            FileType.IMAGE: self.extract_image,
        }

        if file_type in extraction_methods:
            return extraction_methods[file_type](file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def get_file_type(self, file: UploadFile) -> FileType:
        """Determine file type based on filename and content type."""
        if not file.filename:
            return FileType.UNKNOWN

        _, ext = os.path.splitext(file.filename.lower())

        type_mapping = {
            '.pdf': FileType.PDF,
            '.docx': FileType.DOCX,
            '.xlsx': FileType.XLSX,
            '.xls': FileType.XLSX,
            '.jpg': FileType.IMAGE,
            '.jpeg': FileType.IMAGE,
            '.png': FileType.IMAGE,
            '.gif': FileType.IMAGE,
            '.bmp': FileType.IMAGE,
            '.tiff': FileType.IMAGE,
        }

        return type_mapping.get(ext, FileType.UNKNOWN)

    def extract_docx(self, file: UploadFile) -> str:
        """Extract from a DOCX file."""
        try:
            file.file.seek(0)
            document = docx.Document(file.file)
            content = []

            # Extract paragraphs and tables in document order
            for element in document.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find corresponding paragraph object
                    for para in document.paragraphs:
                        if para._element == element:
                            text = para.text.strip()
                            if text:
                                content.append(text)
                            break
                elif element.tag.endswith('tbl'):  # Table
                    # Find corresponding table object
                    for table in document.tables:
                        if table._element == element:
                            md_table = self.__docx_table_to_markdown(table)
                            if md_table:
                                content.append(md_table)
                            break

            # Handle images (stub)
            for rel in document.part.rels.values():
                if 'image' in rel.target_ref:
                    content.append(f"[IMAGE: {os.path.basename(rel.target_ref)}] (OCR not yet implemented)")

            return '\n\n'.join(content)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")

    def __docx_table_to_markdown(self, table) -> str:
        """Convert DOCX table to Markdown format."""
        if not table.rows:
            return ''

        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Clean cell text and handle None values
                cell_text = (cell.text or '').strip().replace('\n', ' ')
                row_data.append(cell_text)
            rows_data.append(row_data)

        if not rows_data:
            return ''

        # Create markdown table
        header = '| ' + ' | '.join(rows_data[0]) + ' |'
        separator = '| ' + ' | '.join(['---'] * len(rows_data[0])) + ' |'
        body_rows = ['| ' + ' | '.join(row) + ' |' for row in rows_data[1:]]

        return '\n'.join([header, separator] + body_rows)

    def __convert_pdf_to_images(self, file_byte: bytes) -> List[Image.Image]:
        """Convert PDF file bytes to a list of PIL Image objects."""
        try:
            images = convert_from_bytes(file_byte)
            return images
        except Exception as e:
            raise ValueError(f"Failed to convert PDF to images: {str(e)}")

    def extract_pdf(self, file: UploadFile) -> str:
        """Extract from a PDF file."""
        try:
            file.file.seek(0)
            file_byte = file.file.read()
            with pdfplumber.open(file.file) as pdf:
                content = []

                for i, page in enumerate(pdf.pages):
                    filtered_page = page
                    chars = filtered_page.chars
                    for table in page.find_tables():
                        first_table_char = page.crop(table.bbox).chars[0]
                        filtered_page = filtered_page.filter(
                            lambda obj:
                                get_bbox_overlap(obj_to_bbox(obj), table.bbox) is None,
                        )
                        chars = filtered_page.chars
                        markdown = self.__table_to_markdown(table.extract())
                        chars.append(first_table_char | {'text': markdown})

                    page_text = extract_text(chars, layout=True)

                    if not page_text:
                        page_text = ''
                        images = self.__convert_pdf_to_images(file_byte)
                        if images:
                            for img in images:
                                page_text += self.__ocr_image(img) + '\n'

                    content.append(page_text)

                return '\n'.join(content)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF file: {str(e)}")

    def __table_to_markdown(self, table: List[List[str]]) -> str:
        """Convert table to Markdown format."""
        if not table or not table[0]:
            return ''

        # Clean table data
        cleaned_table = []
        for row in table:
            cleaned_row = []
            for cell in row:
                if cell is None:
                    cleaned_row.append('')
                else:
                    # Clean cell content
                    cleaned_cell = str(cell).strip().replace('\n', ' ')
                    cleaned_row.append(cleaned_cell)
            cleaned_table.append(cleaned_row)

        # Ensure all rows have the same number of columns
        max_cols = max(len(row) for row in cleaned_table)
        for row in cleaned_table:
            while len(row) < max_cols:
                row.append('')

        # Create markdown table
        header = '| ' + ' | '.join(cleaned_table[0]) + ' |'
        separator = '| ' + ' | '.join(['---'] * max_cols) + ' |'
        body_rows = ['| ' + ' | '.join(row) + ' |' for row in cleaned_table[1:]]

        return '\n'.join([header, separator] + body_rows)

    def __preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """
        Preprocess the image to improve OCR accuracy.
        """
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

        # Apply thresholding to preprocess the image
        thresh = cv2.threshold(
            gray, 0, 255, cv2.THRESH_BINARY
            | cv2.THRESH_OTSU,
        )[1]

        # Apply dilation to merge letters
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilate = cv2.dilate(thresh, kernel, iterations=1)

        return dilate

    def __detect_table_borders(self, preprocessed: np.ndarray):
        # Detect table borders in the preprocessed image.
        horiz_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
        vert_kernel  = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))

        horizontal = cv2.erode(preprocessed, horiz_kernel, iterations=1)
        horizontal = cv2.dilate(horizontal, horiz_kernel, iterations=2)

        vertical = cv2.erode(preprocessed, vert_kernel, iterations=1)
        vertical = cv2.dilate(vertical, vert_kernel, iterations=2)

        table_mask = cv2.add(horizontal, vertical)

        # Find contours of the table borders
        contours, _ = cv2.findContours(
            table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE,
        )
        bboxes = [cv2.boundingRect(c) for c in contours]

        return bboxes

    def __ocr_image(self, image: Image) -> str:
        try:
            np_img = np.array(image)

            preprocessed = self.__preprocess_image(np_img)

            table_bboxes = self.__detect_table_borders(preprocessed)

            tables_markdown = []

            for bbox in table_bboxes:
                x, y, w, h = bbox
                table_crop = image.crop((x, y, x + w, y + h))

                custom_config = r'--oem 3 --psm 6 -l vie'
                table_text = pytesseract.image_to_string(table_crop, config=custom_config)
                table_text = table_text.strip()

                # If Vietnamese fails, try English only
                if not table_text:
                    custom_config = r'--oem 3 --psm 6 -l eng'
                    table_text = pytesseract.image_to_string(table_crop, config=custom_config)
                    table_text = table_text.strip()

                rows = [row.strip().split() for row in table_text.strip().split('\n') if row.strip()]
                markdown = self.__table_to_markdown(rows)
                tables_markdown.append(markdown)

            mask = Image.new('L', image.size, 255)
            draw = ImageDraw.Draw(mask)
            for x, y, w, h in table_bboxes:
                draw.rectangle([x, y, x + w, y + h], fill=0)
            text_only_image = Image.composite(image, Image.new('RGB', image.size, (255, 255, 255)), mask)

            custom_config = r'--oem 3 --psm 6 -l vie'
            text_content = pytesseract.image_to_string(text_only_image, config=custom_config)
            text_content = table_text.strip()

            # If Vietnamese fails, try English only
            if not text_content:
                custom_config = r'--oem 3 --psm 6 -l eng'
                text_content = pytesseract.image_to_string(text_only_image, config=custom_config)
                text_content = table_text.strip()

            text_blocks = []

            if text_content.strip():
                text_blocks.append((0, text_content.strip()))

            for md, bbox in zip(tables_markdown, table_bboxes):
                y = bbox[1]
                if md.strip():
                    text_blocks.append((y, md.strip()))

            text_blocks.sort(key=lambda x: x[0])

            final_content = '\n\n'.join([block[1] for block in text_blocks])
            return final_content

        except Exception as e:
            raise ValueError(f"Failed to extract text from image: {str(e)}")

    def extract_image(self, file: UploadFile) -> str:
        """Extract from an image file using OCR."""
        try:
            file.file.seek(0)
            image = Image.open(file.file)

            content = self.__ocr_image(image)

            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.2)

            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.1)

            custom_config = r'--oem 3 --psm 6 -l vie'
            content = pytesseract.image_to_string(image, config=custom_config)
            content = content.strip()

            # If Vietnamese fails, try English only
            if not content or len(content) < 5:
                custom_config = r'--oem 3 --psm 6 -l eng'
                content = pytesseract.image_to_string(image, config=custom_config)
                content = content.strip()

            return content

        except Exception as e:
            raise ValueError(f"Failed to extract text from image file: {str(e)}")

    def extract_xlsx(self, file: UploadFile) -> str:
        """Extract raw text from an XLSX file and convert tables to Markdown."""
        try:
            file.file.seek(0)
            workbook = openpyxl.load_workbook(BytesIO(file.file.read()))
            content = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Add sheet name as header
                content.append(f"## Sheet: {sheet_name}")

                # Find the actual data range
                max_row = 0
                max_col = 0
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            max_row = max(max_row, cell.row)
                            max_col = max(max_col, cell.column)

                if max_row > 0 and max_col > 0:
                    # Extract data as table
                    table_data = []
                    for row in sheet.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
                        row_data = []
                        for cell in row:
                            value = cell.value
                            if value is None:
                                row_data.append('')
                            else:
                                row_data.append(str(value).strip())
                        table_data.append(row_data)

                    # Convert to markdown table
                    if table_data:
                        md_table = self.__table_to_markdown(table_data)
                        if md_table:
                            content.append(md_table)
                else:
                    content.append('(Empty sheet)')

            return '\n\n'.join(content)
        except Exception as e:
            raise ValueError(f"Failed to extract text from XLSX file: {str(e)}")
